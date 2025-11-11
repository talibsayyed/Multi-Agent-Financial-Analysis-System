"""
File Parsers for Multiple Formats
Handles CSV, Excel, PDF, and DOCX files
"""
import pandas as pd
import PyPDF2
import pdfplumber
from docx import Document
import json
from typing import Dict, Any, List
from datetime import datetime


class BaseParser:
    """Base class for all file parsers"""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """Parse file and return standardized data structure"""
        raise NotImplementedError("Subclasses must implement parse()")
    
    def standardize(self, data: Any) -> Dict[str, Any]:
        """Standardize parsed data to common format"""
        return {
            "parsed_data": data,
            "metadata": {
                "parser": self.__class__.__name__,
                "timestamp": datetime.now().isoformat()
            }
        }


class CSVParser(BaseParser):
    """Parse CSV files into structured financial data"""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse CSV file
        
        Args:
            file_path: Path to CSV file
            
        Returns:
            Standardized data dictionary
        """
        try:
            # Read CSV with pandas
            df = pd.read_csv(file_path)
            
            # Basic data cleaning
            df = self._clean_data(df)
            
            # Extract metadata
            metadata = self._extract_metadata(df, file_path)
            
            # Standardize format
            result = {
                "parsed_data": df,
                "metadata": metadata,
                "columns": list(df.columns),
                "total_records": len(df),
                "summary": self._generate_summary(df)
            }
            
            return result
            
        except Exception as e:
            raise Exception(f"CSV parsing failed: {str(e)}")
    
    def _clean_data(self, df: pd.DataFrame) -> pd.DataFrame:
        """Clean and preprocess CSV data"""
        # Remove empty rows
        df = df.dropna(how='all')
        
        # Convert date columns
        for col in df.columns:
            if 'date' in col.lower():
                try:
                    df[col] = pd.to_datetime(df[col])
                except:
                    pass
        
        # Convert numeric columns
        for col in df.columns:
            if df[col].dtype == 'object':
                try:
                    df[col] = pd.to_numeric(df[col].str.replace(',', ''))
                except:
                    pass
        
        return df
    
    def _extract_metadata(self, df: pd.DataFrame, file_path: str) -> Dict[str, Any]:
        """Extract metadata from CSV data"""
        metadata = {
            "source": file_path,
            "parser": "CSVParser",
            "timestamp": datetime.now().isoformat(),
            "rows": len(df),
            "columns": len(df.columns)
        }
        
        # Detect date range if date column exists
        date_cols = [col for col in df.columns if 'date' in col.lower()]
        if date_cols:
            date_col = date_cols[0]
            metadata["date_range"] = f"{df[date_col].min()} to {df[date_col].max()}"
        
        return metadata
    
    def _generate_summary(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Generate statistical summary of CSV data"""
        summary = {}
        
        numeric_cols = df.select_dtypes(include=['number']).columns
        for col in numeric_cols:
            summary[col] = {
                "mean": float(df[col].mean()),
                "sum": float(df[col].sum()),
                "min": float(df[col].min()),
                "max": float(df[col].max())
            }
        
        return summary


class ExcelParser(BaseParser):
    """Parse Excel files with multiple sheets"""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse Excel file including multiple sheets
        
        Args:
            file_path: Path to Excel file
            
        Returns:
            Standardized data dictionary
        """
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheets_data = {}
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                df = self._clean_data(df)
                sheets_data[sheet_name] = df
            
            # Use first sheet as primary data
            primary_sheet = excel_file.sheet_names[0]
            primary_df = sheets_data[primary_sheet]
            
            metadata = self._extract_metadata(sheets_data, file_path)
            
            result = {
                "parsed_data": primary_df,
                "all_sheets": sheets_data,
                "metadata": metadata,
                "sheet_names": excel_file.sheet_names,
                "columns": list(primary_df.columns),
                "total_records": len(primary_df),
                "summary": self._generate_summary(primary_df)
            }
            
            return result
            
        except Exception as e:
            raise Exception(f"Excel parsing failed: {str(e)}")
    
    def _clean_data(self, df: pd.DataFrame) -> pd.DataFrame:
        """Clean and preprocess Excel data"""
        # Remove completely empty rows and columns
        df = df.dropna(how='all').dropna(axis=1, how='all')
        
        # Convert date columns
        for col in df.columns:
            if 'date' in str(col).lower():
                try:
                    df[col] = pd.to_datetime(df[col])
                except:
                    pass
        
        return df
    
    def _extract_metadata(self, sheets_data: Dict, file_path: str) -> Dict[str, Any]:
        """Extract metadata from Excel file"""
        metadata = {
            "source": file_path,
            "parser": "ExcelParser",
            "timestamp": datetime.now().isoformat(),
            "total_sheets": len(sheets_data),
            "sheet_names": list(sheets_data.keys())
        }
        
        return metadata
    
    def _generate_summary(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Generate summary statistics"""
        summary = {}
        numeric_cols = df.select_dtypes(include=['number']).columns
        
        for col in numeric_cols:
            summary[col] = {
                "mean": float(df[col].mean()),
                "sum": float(df[col].sum())
            }
        
        return summary


class PDFParser(BaseParser):
    """Parse PDF files and extract text and tables"""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse PDF file and extract content
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Standardized data dictionary
        """
        try:
            # Extract text
            text_content = self._extract_text(file_path)
            
            # Extract tables
            tables = self._extract_tables(file_path)
            
            # Parse financial data from text
            financial_data = self._parse_financial_data(text_content, tables)
            
            metadata = {
                "source": file_path,
                "parser": "PDFParser",
                "timestamp": datetime.now().isoformat(),
                "pages": self._get_page_count(file_path),
                "has_tables": len(tables) > 0
            }
            
            result = {
                "parsed_data": financial_data,
                "text_content": text_content[:1000],  # First 1000 chars
                "tables": tables,
                "metadata": metadata
            }
            
            return result
            
        except Exception as e:
            raise Exception(f"PDF parsing failed: {str(e)}")
    
    def _extract_text(self, file_path: str) -> str:
        """Extract text from PDF"""
        text = ""
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text()
        except:
            # Fallback to pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text()
        
        return text
    
    def _extract_tables(self, file_path: str) -> List[pd.DataFrame]:
        """Extract tables from PDF"""
        tables = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_tables = page.extract_tables()
                    for table in page_tables:
                        if table:
                            df = pd.DataFrame(table[1:], columns=table[0])
                            tables.append(df)
        except Exception as e:
            print(f"Table extraction warning: {str(e)}")
        
        return tables
    
    def _parse_financial_data(
        self, 
        text: str, 
        tables: List[pd.DataFrame]
    ) -> Dict[str, Any]:
        """Parse financial data from text and tables"""
        financial_data = {}
        
        # If tables exist, use first table as primary data
        if tables:
            primary_table = tables[0]
            # Try to convert to numeric
            for col in primary_table.columns:
                try:
                    primary_table[col] = pd.to_numeric(
                        primary_table[col].str.replace(',', '').str.replace('$', '')
                    )
                except:
                    pass
            
            financial_data = primary_table.to_dict()
        else:
            # Extract key financial terms from text
            financial_data = self._extract_key_metrics(text)
        
        return financial_data
    
    def _extract_key_metrics(self, text: str) -> Dict[str, Any]:
        """Extract key financial metrics from text"""
        metrics = {}
        
        # Simple pattern matching for common financial terms
        import re
        
        patterns = {
            'revenue': r'revenue[:\s]+\$?([\d,]+)',
            'profit': r'profit[:\s]+\$?([\d,]+)',
            'expenses': r'expenses[:\s]+\$?([\d,]+)'
        }
        
        for key, pattern in patterns.items():
            match = re.search(pattern, text.lower())
            if match:
                try:
                    metrics[key] = float(match.group(1).replace(',', ''))
                except:
                    pass
        
        return metrics
    
    def _get_page_count(self, file_path: str) -> int:
        """Get number of pages in PDF"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                return len(pdf_reader.pages)
        except:
            return 0


class DOCXParser(BaseParser):
    """Parse DOCX files and extract content"""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Standardized data dictionary
        """
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            text_content = self._extract_text(doc)
            
            # Extract tables
            tables = self._extract_tables(doc)
            
            # Parse financial data
            financial_data = self._parse_financial_data(text_content, tables)
            
            metadata = {
                "source": file_path,
                "parser": "DOCXParser",
                "timestamp": datetime.now().isoformat(),
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables)
            }
            
            result = {
                "parsed_data": financial_data,
                "text_content": text_content[:1000],
                "tables": tables,
                "metadata": metadata
            }
            
            return result
            
        except Exception as e:
            raise Exception(f"DOCX parsing failed: {str(e)}")
    
    def _extract_text(self, doc: Document) -> str:
        """Extract text from DOCX"""
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        return '\n'.join(text)
    
    def _extract_tables(self, doc: Document) -> List[pd.DataFrame]:
        """Extract tables from DOCX"""
        tables = []
        
        for table in doc.tables:
            data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                data.append(row_data)
            
            if data:
                df = pd.DataFrame(data[1:], columns=data[0])
                tables.append(df)
        
        return tables
    
    def _parse_financial_data(
        self, 
        text: str, 
        tables: List[pd.DataFrame]
    ) -> Dict[str, Any]:
        """Parse financial data from DOCX content"""
        financial_data = {}
        
        if tables:
            # Use first table
            primary_table = tables[0]
            financial_data = primary_table.to_dict()
        else:
            # Extract from text
            financial_data = self._extract_key_metrics(text)
        
        return financial_data
    
    def _extract_key_metrics(self, text: str) -> Dict[str, Any]:
        """Extract financial metrics from text"""
        metrics = {}
        
        import re
        patterns = {
            'revenue': r'revenue[:\s]+\$?([\d,]+)',
            'profit': r'profit[:\s]+\$?([\d,]+)'
        }
        
        for key, pattern in patterns.items():
            match = re.search(pattern, text.lower())
            if match:
                try:
                    metrics[key] = float(match.group(1).replace(',', ''))
                except:
                    pass
        
        return metrics


class ParserFactory:
    """Factory class to get appropriate parser based on file extension"""
    
    @staticmethod
    def get_parser(file_path: str) -> BaseParser:
        """
        Get appropriate parser for file
        
        Args:
            file_path: Path to file
            
        Returns:
            Parser instance
        """
        extension = file_path.lower().split('.')[-1]
        
        parsers = {
            'csv': CSVParser,
            'xlsx': ExcelParser,
            'xls': ExcelParser,
            'pdf': PDFParser,
            'docx': DOCXParser
        }
        
        parser_class = parsers.get(extension)
        if not parser_class:
            raise ValueError(f"Unsupported file format: {extension}")
        
        return parser_class()
    
    @staticmethod
    def parse_file(file_path: str) -> Dict[str, Any]:
        """
        Parse file using appropriate parser
        
        Args:
            file_path: Path to file
            
        Returns:
            Parsed and standardized data
        """
        parser = ParserFactory.get_parser(file_path)
        return parser.parse(file_path)